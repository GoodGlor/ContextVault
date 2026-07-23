"""Extract text — with positions — from uploaded documents.

The first stage of the ingestion pipeline (design spec §7): turn a raw upload
(PDF / DOCX / TXT) into a ``ParsedDocument`` — the full text plus positioned
``TextBlock``s. Positions are what later let a citation jump to the exact source
passage: every block records its character span into the full text, and (for
PDFs) the 1-based page it came from.

The blocks tile the full text exactly — ``text == "".join(b.text for b in
blocks)`` — so a character offset from a chunk can be mapped back to its block
(and page) by a simple range check.
"""

from dataclasses import dataclass
from io import BytesIO


class DocumentError(Exception):
    """Base class for document-parsing failures."""


class UnsupportedDocumentError(DocumentError):
    """The file's type is not one we can parse."""


class DocumentParseError(DocumentError):
    """The file is of a supported type but could not be read (corrupt/invalid)."""


@dataclass(frozen=True)
class TextBlock:
    """A positioned span of extracted text.

    ``start``/``end`` are character offsets into the parent ``ParsedDocument.text``;
    ``page`` is the 1-based source page when the format has pages (PDF), else None.
    """

    text: str
    start: int
    end: int
    page: int | None


@dataclass(frozen=True)
class ParsedDocument:
    """Full extracted text plus the positioned blocks it is composed of."""

    text: str
    blocks: tuple[TextBlock, ...]


def _blocks_from_segments(segments: list[tuple[str, int | None]]) -> ParsedDocument:
    """Assemble a ``ParsedDocument`` from ``(text, page)`` segments in order."""
    blocks: list[TextBlock] = []
    cursor = 0
    for text, page in segments:
        end = cursor + len(text)
        blocks.append(TextBlock(text=text, start=cursor, end=end, page=page))
        cursor = end
    return ParsedDocument(text="".join(b.text for b in blocks), blocks=tuple(blocks))


def parsed_from_text(text: str) -> ParsedDocument:
    """Wrap ready-made text (e.g. an extracted web page) as a single page-less block."""
    return _blocks_from_segments([(text, None)])


def _parse_txt(data: bytes) -> ParsedDocument:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise DocumentParseError("Text file is not valid UTF-8.") from exc
    return _blocks_from_segments([(text, None)])


def _parse_docx(data: bytes) -> ParsedDocument:
    from docx import Document

    try:
        document = Document(BytesIO(data))
    except Exception as exc:  # python-docx raises package/XML-specific errors
        raise DocumentParseError("Could not read DOCX file.") from exc
    # DOCX has no fixed pagination; each paragraph becomes a page-less block,
    # newline-terminated so blocks tile the full text.
    segments: list[tuple[str, int | None]] = [(p.text + "\n", None) for p in document.paragraphs]
    return _blocks_from_segments(segments)


def _parse_pdf(data: bytes) -> ParsedDocument:
    from pypdf import PdfReader
    from pypdf.errors import PyPdfError

    try:
        reader = PdfReader(BytesIO(data))
        segments: list[tuple[str, int | None]] = [
            (page.extract_text() + "\n", number)
            for number, page in enumerate(reader.pages, start=1)
        ]
    except (PyPdfError, OSError, ValueError) as exc:
        raise DocumentParseError("Could not read PDF file.") from exc
    return _blocks_from_segments(segments)


# Extensions that count as an "image" upload — the single source of truth, also
# consumed by the sources API (to classify a Source's kind) and the ingestion layer
# (to route images to LLM-based OCR instead of ``parse_document``). Images are not in
# ``_PARSERS``: they are transcribed by the repository's vision model, not parsed here.
IMAGE_SUFFIXES: frozenset[str] = frozenset(
    {".png", ".jpg", ".jpeg", ".webp", ".tiff", ".bmp", ".heic", ".heif"}
)

_PARSERS = {
    ".txt": _parse_txt,
    ".docx": _parse_docx,
    ".pdf": _parse_pdf,
}


def file_suffix(filename: str) -> str:
    """Return ``filename``'s lowercased extension including the dot, or ``""``."""
    return "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def parse_document(filename: str, data: bytes) -> ParsedDocument:
    """Parse ``data`` into a ``ParsedDocument``, dispatching on ``filename``'s suffix.

    Raises ``UnsupportedDocumentError`` for unknown types and
    ``DocumentParseError`` for a supported type that cannot be read.
    """
    suffix = file_suffix(filename)
    parser = _PARSERS.get(suffix)
    if parser is None:
        raise UnsupportedDocumentError(
            f"Unsupported document type {suffix or filename!r}; "
            f"supported: {', '.join(sorted(_PARSERS))}."
        )
    return parser(data)
